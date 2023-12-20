import os.path

from flask import request, send_file

from docx import Document

import models


def get_document():
    history_id = int(request.args.get('history_id'))

    db_document = models.Documents.query.filter_by(history_id=history_id).first()

    if db_document is not None:
        file_path = db_document.file_path

        return send_file(file_path, download_name='Договор об аренде.docx')

    history = models.db.get_or_404(models.History, history_id)

    document = Document()

    document.add_heading('Договор об аренде')

    owner = models.db.get_or_404(models.User, history.owner_id)
    renter = models.db.get_or_404(models.User, history.renter_id)
    housing = models.db.get_or_404(models.Housings, history.housing_id)
    address = models.Addresses.query\
        .join(models.Streets)\
        .join(models.Settlements)\
        .join(models.Countries)\
        .filter(models.Addresses.id == housing.address_id)\
        .with_entities(models.Countries.name.label('country_name'),
                       models.Settlements.name.label('settlement_name'),
                       models.Streets.name.label('street_name'),
                       models.Addresses.house_number,
                       models.Addresses.department_number)\
        .first()

    p = document.add_paragraph(f'Я, {owner.first_name} {owner.second_name}, '
                               f'сдаю в аренду {renter.first_name} {renter.second_name}у жилую недвижимость'
                               f', находящеюся по адресу: {address.country_name}, '
                               f'{address.settlement_name}, {address.street_name}, '
                               f'дом: {address.house_number}')

    if address.department_number is not None:
        p.add_run(f', квартира: {address.department_number}.')
    else:
        p.add_run('.')

    document.add_paragraph(f'Цена аренды: {history.price} рублей.')

    document.add_heading('Срок аренды: ')

    document.add_paragraph(f'Дата начала аренды: {history.rent_start}')
    document.add_paragraph(f'Дата окончания аренды: {history.rent_end}')

    document.add_heading(f'Подписи: ')

    document.add_paragraph('Арендодатель: ____________________')
    document.add_paragraph('Дата подписи: ____________________')

    document.add_paragraph('Съемщик: ____________________')
    document.add_paragraph('Дата подписи: ____________________')

    document.add_page_break()

    count_of_documents = models.Documents.query.count()

    directory_path = 'documents/'

    if not os.path.exists(directory_path):
        os.makedirs(directory_path)

    file_path = os.path.join(directory_path, f'doc-{count_of_documents}.docx')

    document.save(file_path)

    db_document = models.Documents(history_id=history_id,
                                   file_path=file_path)

    models.db.session.add(db_document)
    models.db.session.commit()

    return send_file(file_path, download_name='Договор об аренде.docx')
